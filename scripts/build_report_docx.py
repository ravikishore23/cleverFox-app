from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

WORKSPACE = Path(r"c:\Users\ravik\cleverFox-app")
REPORT_DIR = WORKSPACE / "report"
FIG_DIR = REPORT_DIR / "figures"
OUT_PATH = REPORT_DIR / "CleverFox_Project_Report.docx"

# Paths for code excerpts
AGENT_CONTROLLER = WORKSPACE / "backend" / "src" / "agent" / "AgentController.ts"
TASK_MODEL = WORKSPACE / "backend" / "src" / "models" / "Task.ts"
NOTE_ROUTE = WORKSPACE / "backend" / "src" / "routes" / "notes.ts"
AI_CHAT_ROUTE = WORKSPACE / "backend" / "src" / "routes" / "ai" / "chat.ts"
VIDEO_SERVER = WORKSPACE / "backend" / "src" / "realtime" / "videoSessionServer.ts"


def add_centered_title(doc: Document, text: str, size=26):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_centered_text(doc: Document, text: str, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_image(doc: Document, path: Path, width_in=6.0):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
    else:
        add_paragraph(doc, f"[Missing image: {path.name}]")


def add_placeholder_box(doc: Document, label: str, height_lines=6):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = label + "\n" * (height_lines - 1)
    # Add borders
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '777777')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def read_excerpt(path: Path, max_lines=60):
    if not path.exists():
        return "// File not found"
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    return "\n".join(lines[:max_lines])


def main():
    doc = Document()

    # Title page
    add_centered_title(doc, "AI-Based Intelligent Study Environment", size=28)
    add_centered_title(doc, "Using Intelligent Agent Architecture", size=24)
    add_centered_text(doc, "(CleverFox)", size=16, bold=True)
    doc.add_paragraph("\n")
    add_centered_text(doc, "A PROJECT REPORT", size=14, bold=True)
    doc.add_paragraph("\n")
    add_centered_text(doc, "Submitted by", size=12)
    add_centered_text(doc, "<Student Name> (Register No)", size=12)
    add_centered_text(doc, "<Student Name> (Register No)", size=12)
    doc.add_paragraph("\n")
    add_centered_text(doc, "in partial fulfillment for the award of the degree of", size=12)
    add_centered_text(doc, "BACHELOR OF TECHNOLOGY", size=13, bold=True)
    add_centered_text(doc, "ARTIFICIAL INTELLIGENCE AND DATA SCIENCE", size=12, bold=True)
    doc.add_paragraph("\n")
    add_centered_text(doc, "<College Name>", size=12, bold=True)
    add_centered_text(doc, "<University Name>", size=12)
    add_centered_text(doc, "APRIL 2026", size=12)
    doc.add_page_break()

    # Certificate page placeholder
    add_centered_title(doc, "BONAFIDE CERTIFICATE", size=18)
    add_paragraph(doc, "Certified that this project report titled \"AI-Based Intelligent Study Environment Using Intelligent Agent Architecture (CleverFox)\" is the bonafide work carried out by the students under our supervision.")
    doc.add_paragraph("\n\n\n")
    add_paragraph(doc, "Project Guide Signature")
    add_paragraph(doc, "Head of Department Signature")
    doc.add_page_break()

    # Acknowledgement
    add_heading(doc, "Acknowledgement", level=1)
    add_paragraph(doc, "We express our sincere gratitude to our institution and faculty for their guidance, support, and encouragement throughout the development of this project. We thank our project guide for continuous feedback and valuable suggestions. We also thank our peers and family members for their support during the project timeline.")
    doc.add_page_break()

    # Abstract
    add_heading(doc, "Abstract", level=1)
    add_paragraph(doc, "CleverFox is an AI-based intelligent study environment that unifies task management, note organization, scheduling, and AI assistance in a single study room experience. The system integrates a React/Vite frontend, Node/Express backend, MongoDB persistence, and an Electron desktop wrapper. An intelligent agent architecture routes user requests to either conversational AI or a sandboxed local agent capable of controlled file and command operations. The system reduces context switching, improves study efficiency, and enables guided learning through AI-driven feedback and recommendations.")
    doc.add_page_break()

    # Table of contents placeholder
    add_heading(doc, "Table of Contents", level=1)
    add_paragraph(doc, "(Auto-generated by Word if desired)\n")
    add_paragraph(doc, "1. Introduction")
    add_paragraph(doc, "2. Literature Survey")
    add_paragraph(doc, "3. System Analysis")
    add_paragraph(doc, "4. System Architecture")
    add_paragraph(doc, "5. System Design")
    add_paragraph(doc, "6. Technology Stack")
    add_paragraph(doc, "7. Module Description")
    add_paragraph(doc, "8. Implementation")
    add_paragraph(doc, "9. Results and Discussion")
    add_paragraph(doc, "10. Advantages and Limitations")
    add_paragraph(doc, "11. Future Enhancements")
    add_paragraph(doc, "12. Conclusion")
    add_paragraph(doc, "13. References")
    add_paragraph(doc, "14. Appendix")
    doc.add_page_break()

    # List of Figures
    add_heading(doc, "List of Figures", level=1)
    add_paragraph(doc, "Figure 5.3 System Architecture")
    add_paragraph(doc, "Figure 6.1 Data Flow Diagram")
    add_paragraph(doc, "Figure 6.2 Use Case Diagram")
    add_paragraph(doc, "Figure 6.3 Sequence Diagram")
    add_paragraph(doc, "Figure 6.4 Class Diagram")
    doc.add_page_break()

    # Abbreviations
    add_heading(doc, "List of Abbreviations", level=1)
    add_bullets(doc, [
        "AI - Artificial Intelligence",
        "NLP - Natural Language Processing",
        "UI - User Interface",
        "API - Application Programming Interface",
        "DB - Database",
        "MCP - Model Context Protocol",
        "CRUD - Create, Read, Update, Delete",
        "HCI - Human Computer Interaction",
    ])
    doc.add_page_break()

    # Introduction
    add_heading(doc, "Introduction", level=1)
    add_heading(doc, "2.1 Overview of Intelligent Study Systems", level=2)
    add_paragraph(doc, "Intelligent study systems aim to personalize learning by adapting to the student's context, preferences, and progress. These systems combine interactive tools, structured content, and analytics to improve comprehension and retention. CleverFox focuses on integrating the key study utilities (tasks, notes, schedule, and AI assistance) into a single environment.")
    add_paragraph(doc, "In addition to content delivery, intelligent study platforms emphasize workflow efficiency. A unified system lowers context-switching costs, reduces duplicate work, and surfaces relevant study materials at the right time. CleverFox addresses these needs by integrating AI-powered workflows with a focused desktop-first interface.")

    add_heading(doc, "2.2 Importance of AI in Education", level=2)
    add_paragraph(doc, "AI enables personalized tutoring, rapid summarization, and contextual explanations. Natural language interfaces make knowledge retrieval faster and reduce the learning curve for new tools. CleverFox uses AI to provide proactive guidance, summarize content, and respond to study-related questions.")

    add_heading(doc, "2.3 Objectives of the Project", level=2)
    add_bullets(doc, [
        "Provide a single, distraction-minimized study environment for learners.",
        "Integrate AI chat and agent-based assistance into daily study workflows.",
        "Allow structured organization of notes, tasks, and schedule events.",
        "Enable real-time study sessions with collaboration support.",
        "Ensure extensibility through MCP tooling and configurable AI providers.",
    ])
    doc.add_page_break()

    # Literature Survey
    add_heading(doc, "Literature Survey", level=1)
    add_heading(doc, "3.1 Existing Study Platforms", level=2)
    add_paragraph(doc, "Existing platforms provide individual functions: note apps for content storage, task lists for planning, and communication tools for collaboration. These tools often operate independently and require learners to manually connect information across applications. This fragmentation can lower focus and reduce efficiency.")

    add_heading(doc, "3.2 Limitations of Traditional Learning Systems", level=2)
    add_paragraph(doc, "Traditional systems are reactive and require the user to decide what to do next. The lack of automated context leads to repetitive tasks and missed insights. Without AI-driven adaptation, many platforms cannot tailor guidance or surface relevant materials at the right time.")

    add_heading(doc, "3.3 Need for Intelligent Agent-Based Learning", level=2)
    add_paragraph(doc, "Agent-based learning environments can propose actions, perform routine tasks, and summarize learning materials. By bridging the gap between user intent and system execution, agents reduce overhead and allow learners to focus on core learning activities.")
    doc.add_page_break()

    # System Analysis
    add_heading(doc, "System Analysis", level=1)
    add_heading(doc, "4.1 Existing System", level=2)
    add_paragraph(doc, "Students frequently rely on multiple disconnected applications for scheduling, note-taking, and AI queries. This results in fragmented knowledge and repeated context building for each task. Most systems do not retain a unified view of the learner's study context.")

    add_heading(doc, "4.2 Proposed System", level=2)
    add_paragraph(doc, "CleverFox consolidates study tasks into one environment. The backend exposes APIs for notes, tasks, and schedules, while the frontend offers a unified study room interface with AI assistance and real-time video support. A routing layer decides whether to invoke AI chat or a local agent for tool execution.")

    add_heading(doc, "4.3 Advantages of Proposed System", level=2)
    add_bullets(doc, [
        "Single workspace for scheduling, task tracking, note taking, and AI chat.",
        "Reduced cognitive load and fewer context switches.",
        "AI-assisted workflows for explanations, summaries, and suggestions.",
        "Optional local agent for safe automation of study tasks.",
        "Cross-platform compatibility via Electron desktop wrapper.",
    ])
    doc.add_page_break()

    # System Architecture
    add_heading(doc, "System Architecture", level=1)
    add_heading(doc, "5.1 Overview of Intelligent Agent Architecture", level=2)
    add_paragraph(doc, "The architecture is layered: a React-based UI sends requests to an Express backend, which routes AI requests to either a chat provider or a local agent loop. The MCP tool server exposes controlled tools for file and command operations. MongoDB stores persistent data for tasks, notes, schedules, and chats.")

    add_heading(doc, "5.2 Types of Agents Used (Learning Agent, Recommendation Agent, etc.)", level=2)
    add_paragraph(doc, "The Learning Agent responds to user prompts using large language models. The Action Agent (local agent) executes tool actions such as file generation, reading, and running commands. Recommendation behaviors are derived from schedule and task contexts.")

    add_heading(doc, "5.3 Architecture Diagram", level=2)
    add_image(doc, FIG_DIR / "architecture.png", width_in=6.3)
    add_caption(doc, "Figure 5.3: CleverFox architecture overview")
    doc.add_page_break()

    # System Design
    add_heading(doc, "System Design", level=1)
    add_heading(doc, "6.1 Data Flow Diagram (DFD)", level=2)
    add_image(doc, FIG_DIR / "dfd.png", width_in=6.3)
    add_caption(doc, "Figure 6.1: Data flow diagram")

    add_heading(doc, "6.2 Use Case Diagram", level=2)
    add_image(doc, FIG_DIR / "use-case.png", width_in=6.3)
    add_caption(doc, "Figure 6.2: Use case diagram")

    add_heading(doc, "6.3 Sequence Diagram", level=2)
    add_image(doc, FIG_DIR / "sequence.png", width_in=6.3)
    add_caption(doc, "Figure 6.3: Sequence diagram for AI chat")

    add_heading(doc, "6.4 Class Diagram", level=2)
    add_image(doc, FIG_DIR / "class-diagram.png", width_in=6.3)
    add_caption(doc, "Figure 6.4: Core class diagram")
    doc.add_page_break()

    # Technology Stack
    add_heading(doc, "Technology Stack", level=1)
    add_heading(doc, "7.1 Programming Languages Used", level=2)
    add_bullets(doc, ["TypeScript", "JavaScript"])

    add_heading(doc, "7.2 Tools and Frameworks", level=2)
    add_bullets(doc, [
        "Frontend: React, Vite, Chakra UI, Mantine, Tailwind CSS",
        "Backend: Node.js, Express, Socket.IO",
        "Desktop: Electron",
        "Database: MongoDB with Mongoose",
    ])

    add_heading(doc, "7.3 AI Techniques Used (Machine Learning, NLP, etc.)", level=2)
    add_bullets(doc, [
        "Natural Language Processing for conversational queries",
        "Tool-augmented reasoning with agent loops",
        "Web search augmentation for current information",
    ])
    doc.add_page_break()

    # Module Description
    add_heading(doc, "Module Description", level=1)
    add_heading(doc, "8.1 User Registration & Login Module", level=2)
    add_paragraph(doc, "The frontend provides login and signup interfaces with user-friendly validation. This module manages identity state in the client and can be extended for server-side authentication.")

    add_heading(doc, "8.2 Personalized Learning Module", level=2)
    add_paragraph(doc, "The Study Room module integrates timer, notes, tasks, and schedule tools. It is designed to keep users focused by minimizing external context switches and offering a single workspace.")

    add_heading(doc, "8.3 Intelligent Recommendation Module", level=2)
    add_paragraph(doc, "The AI chat and routing layer can propose actions such as generating notes or summarizing content. The system can be extended to proactively recommend study tasks based on deadlines and historical progress.")

    add_heading(doc, "8.4 Performance Tracking Module", level=2)
    add_paragraph(doc, "Performance tracking uses task status and progress metrics. Subtasks are aggregated to compute completion rates and enable progress visualization.")

    add_heading(doc, "8.5 Feedback & Assessment Module", level=2)
    add_paragraph(doc, "The feedback module is implemented through Fox AI. It provides clarifications, explanations, and summary notes. The system can store chat history for revision.")
    doc.add_page_break()

    # Implementation
    add_heading(doc, "Implementation", level=1)
    add_heading(doc, "9.1 Development Process", level=2)
    add_numbered(doc, [
        "UI design and layout for landing, login, and study room pages.",
        "Backend API creation for tasks, notes, schedules, and AI chat.",
        "Integration of MongoDB models for persistent storage.",
        "Local agent setup with controlled file and command tools.",
        "Electron integration for desktop packaging.",
        "Realtime signaling for video study sessions.",
    ])

    add_heading(doc, "9.2 Algorithm/Logic Used", level=2)
    add_paragraph(doc, "Key algorithms include the Pomodoro-style focus timer, agent action routing, and task progress calculation from subtasks. The system uses intent classification to route user requests to either AI chat or tool execution.")

    add_heading(doc, "9.3 Screenshots of the System", level=2)
    add_placeholder_box(doc, "Screenshot 1: Landing Page", height_lines=10)
    add_paragraph(doc, "")
    add_placeholder_box(doc, "Screenshot 2: Study Room - Tasks", height_lines=10)
    add_paragraph(doc, "")
    add_placeholder_box(doc, "Screenshot 3: Fox AI Chat", height_lines=10)
    add_paragraph(doc, "")
    add_placeholder_box(doc, "Screenshot 4: Video Session", height_lines=10)
    doc.add_page_break()

    # Results and Discussion
    add_heading(doc, "Results and Discussion", level=1)
    add_heading(doc, "10.1 Output Screens", level=2)
    add_paragraph(doc, "The outputs demonstrate an integrated study environment where tasks, notes, and schedules are managed in a single interface. The AI assistant provides responses in real time, and video sessions allow collaborative study.")

    add_heading(doc, "10.2 Performance Analysis", level=2)
    add_paragraph(doc, "Frontend responsiveness is maintained using modular React components. Backend CRUD operations remain fast due to indexed MongoDB collections. AI latency is dependent on the configured provider and network conditions.")

    add_heading(doc, "10.3 Comparison with Existing Systems", level=2)
    add_paragraph(doc, "Compared to isolated tools, CleverFox provides a unified context for AI assistance. This reduces repetitive input and improves the quality of recommendations and feedback.")
    doc.add_page_break()

    # Advantages and Limitations
    add_heading(doc, "Advantages and Limitations", level=1)
    add_heading(doc, "11.1 Advantages", level=2)
    add_bullets(doc, [
        "Unified workspace improves focus and workflow efficiency.",
        "Agent architecture enables automation of routine tasks.",
        "Desktop packaging supports a distraction-minimized environment.",
    ])

    add_heading(doc, "11.2 Limitations", level=2)
    add_bullets(doc, [
        "Requires API keys for external AI providers.",
        "Local agent must be carefully controlled for security.",
        "Advanced analytics require further development.",
    ])
    doc.add_page_break()

    # Future Enhancements
    add_heading(doc, "Future Enhancements", level=1)
    add_bullets(doc, [
        "Robust user authentication and profile management.",
        "Offline local models for privacy and speed.",
        "Advanced recommendation engine based on study history.",
        "Collaborative study rooms with shared notes and tasks.",
    ])
    doc.add_page_break()

    # Conclusion
    add_heading(doc, "Conclusion", level=1)
    add_paragraph(doc, "CleverFox demonstrates a practical and extensible AI-based study environment that unifies key learning tools with intelligent agent support. The system architecture is modular, supports real-time collaboration, and provides a strong foundation for future enhancements in personalized education.")
    doc.add_page_break()

    # References
    add_heading(doc, "References", level=1)
    add_numbered(doc, [
        "React Documentation - https://react.dev",
        "Vite Documentation - https://vite.dev",
        "Electron Documentation - https://www.electronjs.org/docs",
        "MongoDB Documentation - https://www.mongodb.com/docs",
        "Model Context Protocol - https://modelcontextprotocol.io",
    ])
    doc.add_page_break()

    # Appendix with code snippets
    add_heading(doc, "Appendix", level=1)
    add_heading(doc, "A. Agent Controller (excerpt)", level=2)
    add_code_block(doc, read_excerpt(AGENT_CONTROLLER, max_lines=50))
    doc.add_paragraph("")
    add_heading(doc, "B. Task Model (excerpt)", level=2)
    add_code_block(doc, read_excerpt(TASK_MODEL, max_lines=50))
    doc.add_paragraph("")
    add_heading(doc, "C. Notes API Route (excerpt)", level=2)
    add_code_block(doc, read_excerpt(NOTE_ROUTE, max_lines=50))
    doc.add_paragraph("")
    add_heading(doc, "D. AI Chat Route (excerpt)", level=2)
    add_code_block(doc, read_excerpt(AI_CHAT_ROUTE, max_lines=50))
    doc.add_paragraph("")
    add_heading(doc, "E. Video Session Server (excerpt)", level=2)
    add_code_block(doc, read_excerpt(VIDEO_SERVER, max_lines=50))

    doc.save(str(OUT_PATH))


if __name__ == "__main__":
    main()
